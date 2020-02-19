from docx import Document
import os
import re

file_list = os.listdir('./warn_docx/')

filter = ['     ', '', '——————————————————————————————————', '                              ']

file_dic = {}
for file in file_list:
    content_list = []
    doc = Document('./warn_docx/' + file)
    for para in doc.paragraphs:
        text = para.text
        if text in filter:
            continue
        else:
            content_list.append(text)
    if len(doc.tables) >= 0:
        tables = doc.tables
        for table in tables:
            row_number = len(table.rows)
            col_number = len(table.columns)
            for i in range(row_number):
                for j in range(col_number):
                    content_list.append(table.cell(i, j).text)
    if content_list == []:
        continue
        # print(file)
    file_dic[file] = content_list

info_dic = {}
messgae_key_words = ['取消', '预计', '目前', '已', '报告', '持续', '解除']
for key, value in file_dic.items():
    info = {}
    title = value[0]
    flag = title.find('机场')
    airport = title[:flag] + '机场'
    print(key)
    print(value)
    print(airport)
    info['airport'] = airport.replace('\n', '').replace(' ', '').replace('\t', '')
    for item in value:
        if item.find("发布时间") != -1 and item.find("北京时") != -1:
            pub_time = re.findall('发布时间(.*?)北京时', item)[0]
            pub_time = pub_time.replace("发布时间：", '').replace('北京时', '').replace('(', '').replace('（', '').replace('：', '')
            info['publish_time'] = pub_time.replace('\t', '')
        if item.find('警报发布序号：') != -1:
            item = item.replace('\n', '').replace(' ', '')
            flag = item.find('警报发布序号：')
            sequence_number = item[flag + 7: flag + 9]
            info['number'] = sequence_number.replace('发', '').replace('\t', '')
        for word in messgae_key_words:
            if item.find(word) != -1:
                info['content'] = item.replace('   ', '').replace('\n', '').replace('\t', '')
    print(info)
    info_dic[key] = info
    print('==========================================================')

file = open('./result data/warn.txt', 'w')
for key, value in info_dic.items():
    try:
        file.write(value['airport'] + '\t' + value['publish_time'] + '\t' + value['number'] + '\t' + value['content'] + '\n')
    except:
        print(key)
file.close()